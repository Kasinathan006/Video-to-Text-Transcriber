#!/usr/bin/env python3
"""
faster_whisper_pipeline.py - High-Speed Local GPU Transcription & Word Document Builder

Transcribes video/audio files or segmented chunks locally using faster-whisper on GPU
and outputs a clean, professional Microsoft Word (.docx) document.
"""

import os
import sys
import glob
import time
import argparse
from faster_whisper import WhisperModel
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def log(msg):
    print(msg, flush=True)


def build_docx_from_transcripts(transcripts, output_path, video_title="Start Your AI Business"):
    log(f"\nBuilding Microsoft Word Document: {output_path}...")
    doc = Document()

    # Configure style with Unicode / Complex Script support
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:cs"), "Nirmala UI")

    # Document Header
    title = doc.add_heading(f"{video_title} - Complete Transcription", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for part_name, text in transcripts:
        doc.add_heading(part_name, level=2)
        if text.strip():
            # Split into paragraphs every few sentences or clean blocks
            paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
            for p in paragraphs:
                doc.add_paragraph(p)
        else:
            doc.add_paragraph("[No speech detected]")

    doc.save(output_path)
    log(f"[SUCCESS] Saved Word document to '{output_path}'.")


def main():
    parser = argparse.ArgumentParser(description="Local GPU Transcription Pipeline via faster-whisper")
    parser.add_argument("--chunks-dir", default="audio_chunks_start_ai", help="Directory containing audio chunks")
    parser.add_argument("--output", default="Start_Your_AI_Business_Full_End_to_End_Transcription.docx", help="Output docx filename")
    parser.add_argument("--model", default="large-v3", help="faster-whisper model name")
    parser.add_argument("--compute-type", default="int8_float16", help="Compute type (int8_float16 fits easily in 4GB RTX 3050 VRAM)")
    args = parser.parse_args()

    audio_files = sorted(glob.glob(os.path.join(args.chunks_dir, "*.wav")))
    if not audio_files:
        log(f"[ERROR] No .wav files found in {args.chunks_dir}")
        sys.exit(1)

    log(f"Found {len(audio_files)} audio chunk(s) in '{args.chunks_dir}'.")
    log(f"Loading faster-whisper model '{args.model}' with compute_type='{args.compute_type}' on CUDA...")
    start_time = time.time()
    
    try:
        model = WhisperModel(args.model, device="cuda", compute_type=args.compute_type)
    except Exception as e:
        log(f"[WARN] Failed to load on CUDA ({e}), trying int8...")
        model = WhisperModel(args.model, device="cuda", compute_type="int8")

    log("Model loaded successfully!")

    transcripts = []
    os.makedirs("txt_outputs_start_ai", exist_ok=True)

    for i, file_path in enumerate(audio_files, start=1):
        filename = os.path.basename(file_path)
        part_title = f"Part {i}: {os.path.splitext(filename)[0]}"
        txt_save_path = os.path.join("txt_outputs_start_ai", f"{os.path.splitext(filename)[0]}.txt")

        # Check if already transcribed
        if os.path.exists(txt_save_path):
            log(f"\n[{i}/{len(audio_files)}] Using cached transcript for {filename}...")
            with open(txt_save_path, "r", encoding="utf-8") as f:
                text = f.read()
            transcripts.append((part_title, text))
            continue

        log(f"\n[{i}/{len(audio_files)}] Transcribing {filename}...")
        chunk_start = time.time()
        
        segments, info = model.transcribe(file_path, beam_size=5, language="en")
        log(f"Detected language '{info.language}' with probability {info.language_probability:.2f}")

        text_lines = []
        for seg_idx, segment in enumerate(segments, start=1):
            text_lines.append(segment.text.strip())
            if seg_idx % 25 == 0:
                elapsed_sec = time.time() - chunk_start
                log(f"  ...processed {seg_idx} segments ({segment.end:.1f}s of audio transcribed in {elapsed_sec:.1f}s)")

        full_text = " ".join(text_lines)
        
        # Save cache immediately
        with open(txt_save_path, "w", encoding="utf-8") as f:
            f.write(full_text)

        elapsed = time.time() - chunk_start
        log(f"-> Completed {filename} in {elapsed:.1f}s ({len(full_text)} chars)")
        transcripts.append((part_title, full_text))

    build_docx_from_transcripts(transcripts, args.output)
    total_time = time.time() - start_time
    log(f"\n============================================================")
    log(f"PIPELINE COMPLETED IN {total_time:.1f} seconds!")
    log(f"Word Document: {args.output}")
    log(f"============================================================")


if __name__ == "__main__":
    main()
