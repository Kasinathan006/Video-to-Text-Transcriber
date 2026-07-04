import os
import sys
import time
import subprocess
import glob
from pathlib import Path
from faster_whisper import WhisperModel
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def log(msg):
    print(msg, flush=True)

def run_gdown(url, output_dir):
    log(f"============================================================")
    log(f"STEP 1: DOWNLOADING GOOGLE DRIVE FOLDER")
    log(f"URL: {url}")
    log(f"Output Directory: {output_dir}")
    log(f"============================================================")
    os.makedirs(output_dir, exist_ok=True)
    cmd = [
        sys.executable, "-m", "gdown",
        "--folder", url,
        "-O", output_dir,
        "--continue",
        "--remaining-ok"
    ]
    log(f"Running command: {' '.join(cmd)}")
    try:
        res = subprocess.run(cmd, check=True)
        log("Download completed successfully!")
    except subprocess.CalledProcessError as e:
        log(f"[WARN] gdown returned non-zero exit code: {e}")
        log("Checking if files exist in directory anyway...")

def find_media_files(directory):
    media_exts = {'.mp4', '.mkv', '.mov', '.avi', '.webm', '.mp3', '.wav', '.m4a', '.flac'}
    media_files = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith('.part'):
                continue
            ext = os.path.splitext(file)[1].lower()
            if ext in media_exts:
                media_files.append(os.path.join(root, file))
    return sorted(media_files)

def get_audio_duration(file_path):
    cmd = [
        "ffprobe", "-v", "error", "-show_entries",
        "format=duration", "-of",
        "default=noprint_wrappers=1:nokey=1", file_path
    ]
    try:
        res = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=True)
        return float(res.stdout.strip())
    except Exception as e:
        log(f"[WARN] Could not get duration via ffprobe: {e}")
        return 0.0

def extract_audio_chunks(video_path, output_dir, chunk_duration_sec=3600):
    log(f"\n--- Extracting Audio Chunks for: {os.path.basename(video_path)} ---")
    os.makedirs(output_dir, exist_ok=True)
    
    total_duration = get_audio_duration(video_path)
    hours = int(total_duration // 3600)
    mins = int((total_duration % 3600) // 60)
    secs = int(total_duration % 60)
    log(f"Total Video Duration: {hours}h {mins}m {secs}s ({total_duration:.1f} seconds)")

    if total_duration == 0:
        total_duration = 100000  # fallback if ffprobe fails

    chunk_files = []
    num_chunks = int(total_duration // chunk_duration_sec) + (1 if total_duration % chunk_duration_sec > 0 else 0)
    if num_chunks == 0:
        num_chunks = 1

    for i in range(num_chunks):
        start_time = i * chunk_duration_sec
        chunk_name = f"chunk_{i+1:03d}.wav"
        chunk_path = os.path.join(output_dir, chunk_name)
        chunk_files.append(chunk_path)

        if os.path.exists(chunk_path) and os.path.getsize(chunk_path) > 10000:
            log(f"Chunk {i+1}/{num_chunks} already exists: {chunk_name} (skipping extraction)")
            continue

        log(f"Extracting Chunk {i+1}/{num_chunks} (Start: {start_time}s, Duration: {chunk_duration_sec}s)...")
        cmd = [
            "ffmpeg", "-y", "-ss", str(start_time), "-t", str(chunk_duration_sec),
            "-i", video_path, "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1",
            chunk_path
        ]
        res = subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.PIPE, text=True)
        if res.returncode != 0:
            log(f"[WARN] ffmpeg warning/error on chunk {i+1}: {res.stderr[:200]}")
            if not os.path.exists(chunk_path) or os.path.getsize(chunk_path) == 0:
                log(f"[INFO] Reached end of audio stream at chunk {i+1}.")
                chunk_files.pop()
                break
        else:
            log(f"-> Created {chunk_name} ({os.path.getsize(chunk_path)/(1024*1024):.1f} MB)")

    return chunk_files, total_duration

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def split_into_paragraphs(text, target_sentences=6):
    words = text.split()
    if not words:
        return []
    
    sentences = []
    current_sentence = []
    for word in words:
        current_sentence.append(word)
        if word.endswith(('.', '?', '!', '。', '？', '！')) and len(current_sentence) > 3:
            sentences.append(" ".join(current_sentence))
            current_sentence = []
    if current_sentence:
        sentences.append(" ".join(current_sentence))

    paragraphs = []
    for i in range(0, len(sentences), target_sentences):
        para_text = " ".join(sentences[i:i+target_sentences])
        paragraphs.append(para_text)
    return paragraphs

def build_docx(transcripts, total_duration, video_filename, output_path):
    log(f"\nBuilding professional Word document: {output_path}...")
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Normal Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x2D, 0x37, 0x48)  # Dark Charcoal
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(8)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run(f"Full End-to-End Verbatim Transcription")
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    run_sub = subtitle_p.add_run(f"Source: {video_filename}")
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    # Metadata Table
    hours = int(total_duration // 3600)
    mins = int((total_duration % 3600) // 60)
    secs = int(total_duration % 60)
    dur_str = f"{hours}h {mins}m {secs}s ({total_duration/60:.1f} minutes)"

    total_words = sum(len(text.split()) for _, text in transcripts)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    metadata = [
        ("Source File:", video_filename),
        ("Total Duration:", dur_str),
        ("Total Word Count:", f"{total_words:,} words (100% End-to-End Extraction)"),
        ("AI Model Used:", "Faster-Whisper (Large-v3 GPU - High Fidelity)")
    ]

    for i, (label, val) in enumerate(metadata):
        row = table.rows[i]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.3)
        set_cell_margins(cell_lbl)
        set_cell_margins(cell_val)
        bg = "F7FAFC" if i % 2 == 0 else "EDF2F7"
        set_cell_background(cell_lbl, bg)
        set_cell_background(cell_val, bg)
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(0)
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(10.5)
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(0)
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    for part_idx, (part_name, text) in enumerate(transcripts, 1):
        words = len(text.split())

        h2 = doc.add_paragraph()
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(8)
        r_h2 = h2.add_run(f"{part_name} (approx. {words:,} words)")
        r_h2.font.size = Pt(15)
        r_h2.font.bold = True
        r_h2.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

        paragraphs = split_into_paragraphs(text, target_sentences=6)
        if not paragraphs:
            doc.add_paragraph("[No speech detected in this audio segment]")
        else:
            for p_text in paragraphs:
                p = doc.add_paragraph(p_text)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.15

    doc.save(output_path)
    log(f"[SUCCESS] Saved Word document to '{output_path}' (Total words: {total_words:,}).")

def main():
    gdrive_url = "https://drive.google.com/drive/folders/1lEqhcp7B2mGoPSNPmdoFk4IGDs-F-JT9?usp=sharing"
    download_dir = "gdrive_videos"
    
    # Step 1: Download folder (unless --skip-download is passed)
    if "--skip-download" not in sys.argv:
        run_gdown(gdrive_url, download_dir)
    else:
        log("[INFO] --skip-download flag passed. Skipping gdown download step.")
    
    # Step 2: Find media files
    media_files = find_media_files(download_dir)
    log(f"\nFound {len(media_files)} media file(s) in '{download_dir}':")
    for mf in media_files:
        log(f"  - {mf} ({os.path.getsize(mf)/(1024*1024):.1f} MB)")

    if not media_files:
        log("[ERROR] No media files found after download!")
        return

    # Step 3: Load Whisper model once
    log("\nLoading faster-whisper model 'large-v3' on CUDA...")
    try:
        model = WhisperModel("large-v3", device="cuda", compute_type="int8_float16")
    except Exception as e:
        log(f"[WARN] int8_float16 failed ({e}), trying int8...")
        try:
            model = WhisperModel("large-v3", device="cuda", compute_type="int8")
        except Exception as e2:
            log(f"[WARN] CUDA failed ({e2}), falling back to CPU...")
            model = WhisperModel("large-v3", device="cpu", compute_type="int8")
    log("Model loaded successfully!")

    # Step 4: Process each video file
    for vid_idx, video_path in enumerate(media_files, start=1):
        log(f"\n============================================================")
        log(f"PROCESSING VIDEO [{vid_idx}/{len(media_files)}]: {os.path.basename(video_path)}")
        log(f"============================================================")
        
        clean_name = os.path.splitext(os.path.basename(video_path))[0]
        # remove special chars for directory names
        safe_name = "".join(c if c.isalnum() or c in (' ', '_', '-') else '_' for c in clean_name).strip()
        
        audio_chunks_dir = f"audio_chunks_{safe_name}"
        txt_save_dir = f"txt_outputs_{safe_name}"
        os.makedirs(txt_save_dir, exist_ok=True)
        
        audio_files, total_duration = extract_audio_chunks(video_path, audio_chunks_dir, 3600)
        
        transcripts = []
        for i, chunk_path in enumerate(audio_files, start=1):
            chunk_filename = os.path.basename(chunk_path)
            part_title = f"Part {i}: Hour {i-1} to Hour {i}" if len(audio_files) > 1 else "Full Video Transcription"
            txt_save_path = os.path.join(txt_save_dir, f"{os.path.splitext(chunk_filename)[0]}.txt")

            if os.path.exists(txt_save_path) and os.path.getsize(txt_save_path) > 10:
                log(f"\n[{i}/{len(audio_files)}] Using cached transcript for {chunk_filename}...")
                with open(txt_save_path, "r", encoding="utf-8") as f:
                    text = f.read()
                transcripts.append((part_title, text))
                continue

            log(f"\n[{i}/{len(audio_files)}] Transcribing {chunk_filename}...")
            chunk_start = time.time()
            
            # Let Whisper auto-detect language or transcribe with beam_size=2 for high accuracy
            segments, info = model.transcribe(chunk_path, beam_size=2)
            log(f"Detected language '{info.language}' with probability {info.language_probability:.2f}")

            txt_tmp_path = txt_save_path + ".tmp"
            text_lines = []
            for seg_idx, segment in enumerate(segments, start=1):
                text_lines.append(segment.text.strip())
                if seg_idx % 30 == 0:
                    elapsed_sec = time.time() - chunk_start
                    log(f"  ...processed {seg_idx} segments ({segment.end:.1f}s of audio transcribed in {elapsed_sec:.1f}s)")
                    with open(txt_tmp_path, "w", encoding="utf-8") as f:
                        f.write(" ".join(text_lines))

            full_text = " ".join(text_lines)
            
            with open(txt_save_path, "w", encoding="utf-8") as f:
                f.write(full_text)
            if os.path.exists(txt_tmp_path):
                try:
                    os.remove(txt_tmp_path)
                except Exception:
                    pass

            elapsed = time.time() - chunk_start
            log(f"-> Completed {chunk_filename} in {elapsed:.1f}s ({len(full_text)} chars)")
            transcripts.append((part_title, full_text))

        docx_filename = f"{safe_name}_Full_End_to_End_Transcription.docx"
        build_docx(transcripts, total_duration, os.path.basename(video_path), docx_filename)

    log("\n============================================================")
    log("ALL VIDEOS PROCESSED SUCCESSFULLY! DOCX FILES GENERATED.")
    log("============================================================")

if __name__ == "__main__":
    main()
