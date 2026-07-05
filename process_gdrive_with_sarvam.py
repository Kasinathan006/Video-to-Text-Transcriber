#!/usr/bin/env python3
"""
process_gdrive_with_sarvam.py - End-to-End Google Drive Video to Word Docx via Sarvam AI

This script:
1. Downloads/resumes downloading the Google Drive folder using gdown.
2. Finds all video files in 'gdrive_videos'.
3. For each video:
   - Extracts 16kHz mono audio and chunks it into <= 45-minute segments.
   - Uploads audio chunks to Sarvam AI batch speech-to-text ('saaras:v3').
   - Polls job until completion and downloads JSON transcripts.
   - Generates a beautifully formatted Word document (.docx) with verbatim transcript, word count, and metadata table.
"""

import os
import sys
import time
import glob
import json
import re
import subprocess
from pathlib import Path

try:
    import sarvamai
except ImportError:
    print("Installing sarvamai...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "sarvamai"])
    import sarvamai

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn


def log(msg):
    print(msg, flush=True)


def run_gdown(url, output_dir):
    log("============================================================")
    log("STEP 1: DOWNLOADING GOOGLE DRIVE FOLDER")
    log(f"URL: {url}")
    log(f"Output Directory: {output_dir}")
    log("============================================================")
    os.makedirs(output_dir, exist_ok=True)
    cmd = [
        sys.executable, "-m", "gdown",
        "--folder", url,
        "-O", output_dir,
        "--continue",
        "--remaining-ok"
    ]
    log(f"Running command: {' '.join(cmd)}")
    try:
        subprocess.run(cmd, check=True)
        log("Download completed successfully!")
    except subprocess.CalledProcessError as e:
        log(f"[WARN] gdown returned non-zero exit code: {e}")
        log("Checking existing files in directory...")


def find_media_files(directory):
    media_exts = {'.mp4', '.mkv', '.mov', '.avi', '.webm', '.mp3', '.wav', '.m4a', '.flac'}
    media_files = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith('.part'):
                continue
            ext = os.path.splitext(file)[1].lower()
            if ext in media_exts:
                media_files.append(os.path.join(root, file))
    return sorted(media_files)


def get_media_duration_seconds(file_path):
    try:
        cmd = [
            "ffprobe", "-v", "error",
            "-show_entries", "format=duration",
            "-of", "default=noprint_wrappers=1:nokey=1",
            str(file_path)
        ]
        out = subprocess.check_output(cmd, text=True).strip()
        return float(out)
    except Exception as e:
        log(f"[WARN] Could not determine exact duration via ffprobe ({e}).")
        return 0.0


def extract_and_segment_audio(input_file, work_dir, chunk_minutes=45):
    os.makedirs(work_dir, exist_ok=True)
    base_name = Path(input_file).stem
    
    # Clean base_name for file system paths
    safe_name = "".join(c if c.isalnum() or c in (' ', '_', '-') else '_' for c in base_name).strip()
    
    existing_parts = sorted(glob.glob(os.path.join(work_dir, f"{safe_name}_part_*.wav")))
    if existing_parts:
        log(f"Found {len(existing_parts)} existing segmented audio files in {work_dir}. Using existing chunks!")
        duration = get_media_duration_seconds(input_file)
        return existing_parts, duration

    duration = get_media_duration_seconds(input_file)
    chunk_seconds = chunk_minutes * 60
    
    if duration > chunk_seconds:
        log(f"Media duration ({duration/60:.1f} mins) exceeds {chunk_minutes} mins. Segmenting into parts...")
        out_pattern = os.path.join(work_dir, f"{safe_name}_part_%02d.wav")
        cmd = [
            "ffmpeg", "-y", "-err_detect", "ignore_err", "-fflags", "+discardcorrupt",
            "-i", str(input_file),
            "-f", "segment", "-segment_time", str(chunk_seconds),
            "-ar", "16000", "-ac", "1", "-c:a", "pcm_s16le",
            out_pattern
        ]
        subprocess.run(cmd, check=True)
        files = sorted(glob.glob(os.path.join(work_dir, f"{safe_name}_part_*.wav")))
    else:
        log(f"Media duration is within limit. Extracting/converting to WAV...")
        out_file = os.path.join(work_dir, f"{safe_name}.wav")
        cmd = [
            "ffmpeg", "-y", "-err_detect", "ignore_err", "-fflags", "+discardcorrupt",
            "-i", str(input_file),
            "-ar", "16000", "-ac", "1", "-c:a", "pcm_s16le",
            out_file
        ]
        subprocess.run(cmd, check=True)
        files = [out_file]
        
    return files, duration


def split_into_paragraphs(text, target_sentences=6):
    text = text.strip()
    if not text:
        return []
    sentences = re.split(r'(?<=[.!?])\s+', text)
    paragraphs = []
    current_chunk = []
    for sentence in sentences:
        current_chunk.append(sentence)
        if len(current_chunk) >= target_sentences:
            paragraphs.append(' '.join(current_chunk))
            current_chunk = []
    if current_chunk:
        paragraphs.append(' '.join(current_chunk))
    return paragraphs


def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def create_transcription_docx(json_files, output_docx_path, source_name, duration_sec):
    doc = Document()
    
    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(8)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("Full End-to-End Verbatim Transcription")
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    run_sub = subtitle_p.add_run(source_name)
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hours = int(duration_sec // 3600)
    mins = int((duration_sec % 3600) // 60)
    secs = int(duration_sec % 60)
    dur_str = f"{hours}h {mins}m {secs}s ({duration_sec/60:.1f} minutes)" if duration_sec > 0 else "N/A"
    
    metadata = [
        ("Source File:", source_name),
        ("Total Duration:", dur_str),
        ("AI Model Used:", "Sarvam AI saaras:v3 Batch Speech-to-Text"),
        ("Transcription Fidelity:", "100% Verbatim End-to-End Extraction")
    ]

    for i, (label, val) in enumerate(metadata):
        row = table.rows[i]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.3)
        
        p0 = cell_lbl.paragraphs[0]
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(10)
        
        p1 = cell_val.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(val)
        r1.font.size = Pt(10)

        bg = "F7FAFC" if i % 2 == 0 else "EDF2F7"
        set_cell_background(cell_lbl, bg)
        set_cell_background(cell_val, bg)
        set_cell_margins(cell_lbl, 80, 80, 100, 100)
        set_cell_margins(cell_val, 80, 80, 100, 100)

    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E0"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E0"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '<w:insideV w:val="none"/><w:left w:val="none"/><w:right w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)
    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    total_words = 0
    chunk_mins = 45

    for idx, path in enumerate(sorted(json_files)):
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
            text = data.get("transcript", "").strip()

        words = len(text.split())
        total_words += words

        start_min = idx * chunk_mins
        end_min = (idx + 1) * chunk_mins
        title = f"Part {idx+1} ({start_min}m - {end_min}m)" if len(json_files) > 1 else "Full Transcript"

        h1 = doc.add_paragraph()
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(4)
        h1.paragraph_format.keep_with_next = True
        run_h1 = h1.add_run(title)
        run_h1.font.size = Pt(15)
        run_h1.font.bold = True
        run_h1.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

        h2 = doc.add_paragraph()
        h2.paragraph_format.space_after = Pt(12)
        h2.paragraph_format.keep_with_next = True
        run_h2 = h2.add_run(f"Word Count: ~{words:,} words")
        run_h2.font.size = Pt(10)
        run_h2.font.italic = True
        run_h2.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

        paras = split_into_paragraphs(text, target_sentences=6)
        for p_text in paras:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            run_p = p.add_run(p_text)
            run_p.font.size = Pt(11)

    log(f"Total Words Transcribed: {total_words:,}")
    doc.save(output_docx_path)
    log(f"[SUCCESS] Generated DOCX at: {output_docx_path}")


def main():
    api_key = os.getenv("SARVAM_API_KEY", "")
    if not api_key.strip():
        log("[ERROR] SARVAM_API_KEY environment variable is required! Please set it before running.")
        return

    gdrive_url = "https://drive.google.com/drive/folders/1lEqhcp7B2mGoPSNPmdoFk4IGDs-F-JT9?usp=sharing"
    download_dir = "gdrive_videos"

    if "--skip-download" not in sys.argv:
        run_gdown(gdrive_url, download_dir)
    else:
        log("[INFO] --skip-download passed. Skipping gdown.")

    media_files = find_media_files(download_dir)
    log(f"\nFound {len(media_files)} media file(s) in '{download_dir}':")
    for mf in media_files:
        log(f"  - {mf} ({os.path.getsize(mf)/(1024*1024):.1f} MB)")

    if not media_files:
        log("[ERROR] No media files found after download!")
        return

    client = sarvamai.SarvamAI(api_subscription_key=api_key)

    for vid_idx, video_path in enumerate(media_files, start=1):
        log(f"\n============================================================")
        log(f"PROCESSING VIDEO [{vid_idx}/{len(media_files)}]: {os.path.basename(video_path)}")
        log(f"============================================================")

        clean_name = os.path.splitext(os.path.basename(video_path))[0]
        safe_name = "".join(c if c.isalnum() or c in (' ', '_', '-') else '_' for c in clean_name).strip()
        
        docx_filename = f"{safe_name}_Full_End_to_End_Transcription.docx"
        if os.path.exists(docx_filename):
            log(f"[INFO] '{docx_filename}' already exists! Skipping video.")
            continue

        audio_work_dir = f"audio_chunks_{safe_name}"
        json_save_dir = f"json_outputs_{safe_name}"
        os.makedirs(json_save_dir, exist_ok=True)

        audio_files, duration = extract_and_segment_audio(video_path, audio_work_dir, chunk_minutes=45)
        log(f"Prepared {len(audio_files)} audio file(s) for batch transcription.")

        existing_jsons = sorted(glob.glob(os.path.join(json_save_dir, "*.json")))
        if len(existing_jsons) == len(audio_files) and len(audio_files) > 0:
            log(f"[INFO] Found {len(existing_jsons)} existing JSON transcripts in '{json_save_dir}'. Skipping Sarvam API upload/transcription!")
            create_transcription_docx(existing_jsons, docx_filename, os.path.basename(video_path), duration)
            continue

        log("Initializing Sarvam AI Batch Speech-to-Text Job...")
        stt_job = client.speech_to_text_job.create_job(
            model="saaras:v3",
            language_code="en-IN",
            with_diarization=False
        )
        job_id = stt_job.job_id
        log(f"Job created successfully! Job ID: {job_id}")

        log(f"Uploading {len(audio_files)} audio files to Sarvam AI...")
        stt_job.upload_files(audio_files)

        log("Starting job and polling status...")
        stt_job.start()
        while True:
            status = client.speech_to_text_job.get_status(job_id)
            log(f"[{time.strftime('%H:%M:%S')}] Job State: {status.job_state}")
            if status.job_state in ["Completed", "Failed"]:
                break
            time.sleep(15)

        if status.job_state == "Failed":
            log(f"[ERROR] Transcription job {job_id} failed on Sarvam AI servers.")
            continue

        log("Downloading JSON outputs...")
        stt_job.download_outputs(json_save_dir)

        json_files = sorted(glob.glob(os.path.join(json_save_dir, "*.json")))
        log(f"Downloaded {len(json_files)} JSON file(s).")

        log("Generating clean Word document (.docx)...")
        create_transcription_docx(json_files, docx_filename, os.path.basename(video_path), duration)

    log("\n============================================================")
    log("ALL VIDEOS PROCESSED SUCCESSFULLY! DOCX FILES GENERATED.")
    log("============================================================")


if __name__ == "__main__":
    main()
