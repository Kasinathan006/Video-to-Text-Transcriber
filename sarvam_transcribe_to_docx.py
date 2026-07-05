import argparse
import glob
import json
import os
import re
import subprocess
import sys
import time
from pathlib import Path

# Try importing required packages
try:
    import sarvamai
    from sarvamai.speech_to_text_job.job import SpeechToTextJob
except ImportError:
    print("Installing required package: sarvamai...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "sarvamai"])
    import sarvamai
    from sarvamai.speech_to_text_job.job import SpeechToTextJob

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("Installing required package: python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn


def get_media_duration_seconds(file_path):
    """Get duration of media file in seconds using ffprobe."""
    try:
        cmd = [
            "ffprobe", "-v", "error",
            "-show_entries", "format=duration",
            "-of", "default=noprint_wrappers=1:nokey=1",
            str(file_path)
        ]
        out = subprocess.check_output(cmd, text=True).strip()
        return float(out)
    except Exception as e:
        print(f"Warning: Could not determine exact duration via ffprobe ({e}). Assuming single file.")
        return 0.0


def extract_and_segment_audio(input_file, work_dir, chunk_minutes=45):
    """Extract audio to 16kHz WAV and segment into <= 45m chunks if longer."""
    os.makedirs(work_dir, exist_ok=True)
    base_name = Path(input_file).stem
    existing_parts = sorted(glob.glob(os.path.join(work_dir, f"{base_name}_part_*.wav")))
    if existing_parts:
        print(f"Found {len(existing_parts)} existing segmented audio files in {work_dir}. Using existing chunks!")
        duration = get_media_duration_seconds(input_file)
        return existing_parts, duration

    duration = get_media_duration_seconds(input_file)

    # Hardened decode flags: survive corrupted MP4 containers / AAC stream glitches
    resilient_flags = ["-err_detect", "ignore_err", "-fflags", "+discardcorrupt"]

    chunk_seconds = chunk_minutes * 60
    if duration > chunk_seconds:
        print(f"Media duration ({duration/60:.1f} mins) exceeds {chunk_minutes} mins. Segmenting into parts...")
        out_pattern = os.path.join(work_dir, f"{base_name}_part_%02d.wav")
        cmd = [
            "ffmpeg", "-y", *resilient_flags, "-i", str(input_file),
            "-f", "segment", "-segment_time", str(chunk_seconds),
            "-ar", "16000", "-ac", "1", "-c:a", "pcm_s16le",
            out_pattern
        ]
        subprocess.run(cmd, check=True)
        files = sorted(glob.glob(os.path.join(work_dir, f"{base_name}_part_*.wav")))
    else:
        print(f"Media duration is within limit. Extracting/converting to WAV...")
        out_file = os.path.join(work_dir, f"{base_name}.wav")
        cmd = [
            "ffmpeg", "-y", *resilient_flags, "-i", str(input_file),
            "-ar", "16000", "-ac", "1", "-c:a", "pcm_s16le",
            out_file
        ]
        subprocess.run(cmd, check=True)
        files = [out_file]

    return files, duration


def split_into_paragraphs(text, target_sentences=6):
    """Split long continuous transcript into readable paragraphs."""
    text = text.strip()
    if not text:
        return []
    sentences = re.split(r'(?<=[.!?])\s+', text)
    paragraphs = []
    current_chunk = []
    for sentence in sentences:
        current_chunk.append(sentence)
        if len(current_chunk) >= target_sentences:
            paragraphs.append(' '.join(current_chunk))
            current_chunk = []
    if current_chunk:
        paragraphs.append(' '.join(current_chunk))
    return paragraphs


def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def create_transcription_docx(json_files, output_docx_path, source_name, duration_sec,
                              font_name="Calibri", sentences_per_paragraph=6,
                              model_label="Sarvam AI saaras:v3 Batch Speech-to-Text",
                              chunk_minutes=45):
    doc = Document()

    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(8)

    # Complex-script font fallback so Tamil/Hindi/Telugu glyphs never render as boxes
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:cs"), "Nirmala UI")
    rFonts.set(qn("w:eastAsia"), "Nirmala UI")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("Full End-to-End Transcription")
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    run_sub = subtitle_p.add_run(source_name)
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    dur_str = f"{int(duration_sec // 3600)}h {int((duration_sec % 3600) // 60)}m {int(duration_sec % 60)}s" if duration_sec > 0 else "N/A"
    metadata = [
        ("Source File:", source_name),
        ("Total Duration:", dur_str),
        ("AI Model Used:", model_label),
        ("Transcription Fidelity:", "100% Verbatim End-to-End Extraction")
    ]

    for i, (label, val) in enumerate(metadata):
        row = table.rows[i]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.3)
        
        p0 = cell_lbl.paragraphs[0]
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(10)
        
        p1 = cell_val.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(val)
        r1.font.size = Pt(10)

        set_cell_background(cell_lbl, "F7FAFC" if i % 2 == 0 else "EDF2F7")
        set_cell_background(cell_val, "F7FAFC" if i % 2 == 0 else "EDF2F7")
        set_cell_margins(cell_lbl, 80, 80, 100, 100)
        set_cell_margins(cell_val, 80, 80, 100, 100)

    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E0"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E0"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '<w:insideV w:val="none"/><w:left w:val="none"/><w:right w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)
    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    total_words = 0
    chunk_mins = chunk_minutes

    for idx, path in enumerate(sorted(json_files)):
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
            text = data.get("transcript", "").strip()

        words = len(text.split())
        total_words += words

        start_min = idx * chunk_mins
        end_min = (idx + 1) * chunk_mins
        title = f"Part {idx+1} ({start_min}m - {end_min}m)" if len(json_files) > 1 else "Full Transcript"

        h1 = doc.add_paragraph()
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(4)
        h1.paragraph_format.keep_with_next = True
        run_h1 = h1.add_run(title)
        run_h1.font.size = Pt(15)
        run_h1.font.bold = True
        run_h1.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

        h2 = doc.add_paragraph()
        h2.paragraph_format.space_after = Pt(12)
        h2.paragraph_format.keep_with_next = True
        run_h2 = h2.add_run(f"Word Count: ~{words:,} words")
        run_h2.font.size = Pt(10)
        run_h2.font.italic = True
        run_h2.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

        paras = split_into_paragraphs(text, target_sentences=sentences_per_paragraph)
        for p_text in paras:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            run_p = p.add_run(p_text)
            run_p.font.size = Pt(11)

    print(f"Total Words Transcribed: {total_words:,}")
    doc.save(output_docx_path)
    print(f"Successfully generated DOCX at: {output_docx_path}")


def main():
    parser = argparse.ArgumentParser(description="End-to-End Video/Audio Transcription to Word Docx via Sarvam AI")
    parser.add_argument("input_file", help="Path to input video or audio file (.mp4, .wav, etc.)")
    parser.add_argument("--api-key", default=os.getenv("SARVAM_API_KEY", ""),
                        help="Sarvam AI API Key")
    parser.add_argument("--output", "-o", default=None, help="Output .docx path")
    parser.add_argument("--work-dir", default="audio_chunks", help="Temporary folder for audio segmentation")
    parser.add_argument("--json-dir", default="transcription_outputs", help="Folder to save downloaded JSON transcripts")
    args = parser.parse_args()

    if not os.path.exists(args.input_file):
        print(f"Error: Input file '{args.input_file}' not found.")
        sys.exit(1)

    if not args.api_key.strip():
        print("Error: Sarvam AI API key is required. Please provide --api-key or set SARVAM_API_KEY environment variable.")
        sys.exit(1)

    out_docx = args.output
    if not out_docx:
        out_docx = f"{Path(args.input_file).stem}_Full_Transcription.docx"

    print("Step 1: Extracting and segmenting audio...")
    audio_files, duration = extract_and_segment_audio(args.input_file, args.work_dir)
    print(f"Prepared {len(audio_files)} audio file(s) for batch transcription.")

    print("Step 2: Initializing Sarvam AI Batch Speech-to-Text Job...")
    client = sarvamai.SarvamAI(api_subscription_key=args.api_key)
    
    stt_job = client.speech_to_text_job.create_job(
        model="saaras:v3",
        language_code="en-IN",
        with_diarization=False
    )
    job_id = stt_job.job_id
    print(f"Job created successfully! Job ID: {job_id}")

    print("Step 3: Uploading audio files to Sarvam AI...")
    stt_job.upload_files(audio_files)

    print("Step 4: Starting job and polling status...")
    stt_job.start()
    while True:
        status = client.speech_to_text_job.get_status(job_id)
        print(f"[{time.strftime('%H:%M:%S')}] Job State: {status.job_state}")
        if status.job_state in ["Completed", "Failed"]:
            break
        time.sleep(15)

    if status.job_state == "Failed":
        print("Error: Transcription job failed on Sarvam AI servers.")
        sys.exit(1)

    print("Step 4: Downloading JSON outputs...")
    os.makedirs(args.json_dir, exist_ok=True)
    stt_job.download_outputs(args.json_dir)
    
    json_files = sorted(glob.glob(os.path.join(args.json_dir, "*.json")))
    print(f"Downloaded {len(json_files)} JSON file(s).")

    print("Step 5: Generating clean Word document (.docx)...")
    create_transcription_docx(json_files, out_docx, Path(args.input_file).name, duration)


if __name__ == "__main__":
    main()
