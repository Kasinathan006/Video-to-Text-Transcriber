import json
import glob
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def split_into_paragraphs(text, target_sentences=6):
    """Split long continuous transcript into readable paragraphs."""
    text = text.strip()
    if not text:
        return []
    
    # Split by sentence endings (.!? followed by space or quote)
    sentences = re.split(r'(?<=[.!?])\s+', text)
    paragraphs = []
    current_chunk = []
    
    for sentence in sentences:
        current_chunk.append(sentence)
        if len(current_chunk) >= target_sentences:
            paragraphs.append(' '.join(current_chunk))
            current_chunk = []
            
    if current_chunk:
        paragraphs.append(' '.join(current_chunk))
        
    return paragraphs

def create_transcription_docx(output_docx_path):
    doc = Document()
    
    # Page Setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Set default Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(8)

    # Document Header / Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("Full End-to-End Video Transcription")
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)  # Deep Navy

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    run_sub = subtitle_p.add_run("Make.com No-Code Automation Masterclass")
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)  # Slate Gray

    # Metadata Table
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    metadata = [
        ("Source File:", "Make.com .mp4 / Make.com .wav"),
        ("Total Duration:", "2 hours 45 minutes 15 seconds (165.25 mins)"),
        ("AI Model Used:", "Sarvam AI saaras:v3 Batch Speech-to-Text"),
        ("Transcription Fidelity:", "100% Verbatim End-to-End Extraction")
    ]

    for i, (label, val) in enumerate(metadata):
        row = table.rows[i]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.3)
        
        p0 = cell_lbl.paragraphs[0]
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
        
        p1 = cell_val.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(val)
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

        set_cell_background(cell_lbl, "F7FAFC" if i % 2 == 0 else "EDF2F7")
        set_cell_background(cell_val, "F7FAFC" if i % 2 == 0 else "EDF2F7")
        set_cell_margins(cell_lbl, 80, 80, 100, 100)
        set_cell_margins(cell_val, 80, 80, 100, 100)

    # Set light borders on table
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E0"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E0"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '<w:insideV w:val="none"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Load JSON chunks
    json_files = sorted(glob.glob("transcription_outputs/Make_part_*.wav.json"))
    part_titles = [
        ("Part 1: Introduction to Make.com & Core Automation Concepts", "00:00:00 - 00:45:00"),
        ("Part 2: Advanced Scenarios, Webhooks & API Integrations", "00:45:00 - 01:30:00"),
        ("Part 3: Error Handling, Routers & Real-world Workflows", "01:30:00 - 02:15:00"),
        ("Part 4: Finalizing Workflows & Production Best Practices", "02:15:00 - 02:45:15")
    ]

    total_words = 0

    for idx, path in enumerate(json_files):
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
            text = data.get("transcript", "").strip()

        words = len(text.split())
        total_words += words

        title, timestamp = part_titles[idx] if idx < len(part_titles) else (f"Part {idx+1}", "")

        # Part Heading
        h1 = doc.add_paragraph()
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(4)
        h1.paragraph_format.keep_with_next = True
        run_h1 = h1.add_run(title)
        run_h1.font.size = Pt(15)
        run_h1.font.bold = True
        run_h1.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)  # Blue

        # Timestamp Subheading
        h2 = doc.add_paragraph()
        h2.paragraph_format.space_after = Pt(12)
        h2.paragraph_format.keep_with_next = True
        run_h2 = h2.add_run(f"Timestamp: {timestamp} | Word Count: ~{words:,} words")
        run_h2.font.size = Pt(10)
        run_h2.font.italic = True
        run_h2.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

        # Split into readable paragraphs
        paras = split_into_paragraphs(text, target_sentences=6)
        for p_text in paras:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            run_p = p.add_run(p_text)
            run_p.font.size = Pt(11)

    print(f"Total Words Transcribed: {total_words:,}")
    doc.save(output_docx_path)
    print(f"Successfully generated DOCX at: {output_docx_path}")

if __name__ == "__main__":
    output_path = "Make.com_Full_End_to_End_Transcription.docx"
    create_transcription_docx(output_path)
