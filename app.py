#!/usr/bin/env python3
"""
app.py - VoxDoc AI / ScribeFlow — AI Media-to-Document Studio (MVP Web Product)

Ingests video/audio via direct upload or Google Drive link, transcribes with
Sarvam AI Batch API (cloud) or Faster-Whisper (local/offline), and generates
professionally formatted Microsoft Word (.docx) documents with optional
speaker diarization.

To launch:
    pip install -r requirements.txt
    streamlit run app.py
"""

import glob
import json
import os
import re
import time
import uuid
from pathlib import Path

import requests
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor

import sarvam_transcribe_to_docx as engine

APP_DIR = Path(__file__).resolve().parent
WORKSPACE_ROOT = APP_DIR / "voxdoc_workspace"

MEDIA_EXTENSIONS = {
    ".mp4", ".mkv", ".mov", ".avi", ".flv", ".wmv", ".webm", ".m4v",
    ".wav", ".mp3", ".m4a", ".flac", ".ogg", ".aac",
}

LANGUAGE_OPTIONS = {
    "English / Tanglish / Multilingual (en-IN)": "en-IN",
    "Hindi (hi-IN)": "hi-IN",
    "Tamil (ta-IN)": "ta-IN",
    "Telugu (te-IN)": "te-IN",
    "Malayalam (ml-IN)": "ml-IN",
    "Kannada (kn-IN)": "kn-IN",
    "Bengali (bn-IN)": "bn-IN",
    "Marathi (mr-IN)": "mr-IN",
}

SARVAM_MODELS = {
    "saaras:v3 — English document output (proven)": "saaras:v3",
    "saarika:v2 — native-language document output": "saarika:v2",
}

FONT_OPTIONS = {
    "Calibri": "Calibri",
    "Arial": "Arial",
    "Nirmala UI (Best for Tamil/Hindi)": "Nirmala UI",
}

st.set_page_config(
    page_title="VoxDoc AI — Video to Formatted Docx Studio",
    page_icon="🎙️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# --- Design System: vibrant dark mode + glassmorphism (Build Guide §5) ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;800&display=swap');
    .stApp {
        background: radial-gradient(1100px 700px at 12% 0%, #1e1b4b 0%, #0B0F19 55%);
        color: #f8fafc;
    }
    .main-header {
        font-family: 'Outfit', 'Inter', sans-serif;
        font-weight: 800;
        font-size: 2.8rem;
        background: linear-gradient(90deg, #818cf8, #c084fc, #f472b6);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 0.2rem;
    }
    .sub-header { font-size: 1.1rem; color: #94a3b8; margin-bottom: 2rem; }
    .glass-card {
        background: rgba(255, 255, 255, 0.03);
        border: 1px solid rgba(255, 255, 255, 0.08);
        border-radius: 16px;
        padding: 1.2rem 1.4rem;
        margin-bottom: 1.2rem;
        backdrop-filter: blur(12px);
    }
    .stButton>button {
        background: linear-gradient(135deg, #6366f1 0%, #8b5cf6 100%);
        color: white; font-weight: 600; border: none; border-radius: 8px;
        padding: 0.6rem 1.5rem; transition: all 0.3s ease;
    }
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 20px rgba(99, 102, 241, 0.4);
    }
    .stProgress > div > div > div > div {
        background-image: linear-gradient(to right, #6366F1, #8B5CF6);
    }
    div[data-testid="stFileUploaderDropzone"] {
        border: 1.5px dashed rgba(139, 92, 246, 0.45) !important;
        background: rgba(99, 102, 241, 0.05) !important;
    }
    div[data-testid="stFileUploaderDropzone"]:hover {
        border-color: #8B5CF6 !important;
        box-shadow: 0 0 22px rgba(139, 92, 246, 0.25);
    }
    div[data-testid="stMetricValue"] { color: #10B981; font-weight: 700; }
</style>
""", unsafe_allow_html=True)


# ----------------------------------------------------------------------------
# Security & Authentication ("naa matum use pannura mathuri" - Personal Mode)
# ----------------------------------------------------------------------------

def get_secret(key: str, default: str = "") -> str:
    """Safely fetch a secret from Streamlit secrets or OS environment variables."""
    val = os.getenv(key, "")
    if not val:
        try:
            if key in st.secrets:
                val = str(st.secrets[key])
        except Exception:
            pass
    return val or default


def check_password():
    """Returns True if the user has entered the correct password."""
    if st.session_state.get("authenticated", False):
        return True

    admin_password = get_secret("VOXDOC_ADMIN_PASSWORD", "") or get_secret("APP_PASSWORD", "") or "mohan"

    st.markdown("<div class='main-header' style='text-align: center; margin-top: 3rem;'>🔒 VoxDoc AI Studio</div>", unsafe_allow_html=True)
    st.markdown("<div class='sub-header' style='text-align: center;'>Private Cloud Workspace — Personal Authentication Required</div>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        with st.container(border=True):
            st.markdown("### 🔑 Enter Personal Passcode")
            st.caption("This workspace is locked for private use (`naa matum use pannura mathuri`). Enter your passcode to access the studio.")
            
            pwd_input = st.text_input("Passcode / PIN", type="password", key="login_pwd_input", placeholder="Enter passcode...")
            
            if st.button("🔓 Unlock Workspace", type="primary", use_container_width=True):
                if pwd_input == admin_password:
                    st.session_state["authenticated"] = True
                    st.success("Access Granted! Welcome to your private cloud studio.")
                    time.sleep(0.5)
                    st.rerun()
                else:
                    st.error("❌ Incorrect Passcode! Access Denied.")
            
            st.markdown("---")
            st.info("💡 **Note for Admin:** Default passcode is `mohan` if not configured. To customize, set `VOXDOC_ADMIN_PASSWORD` in your cloud `.env` or Streamlit Secrets.")
            
    return False


@st.cache_data(show_spinner=False, ttl=300)
def verify_api_key(engine: str, key: str) -> tuple[bool, str]:
    """Verify if the provided API key is valid by testing against cloud API endpoints."""
    if not key or not key.strip():
        return False, "API Key is empty."
    key = key.strip()
    try:
        if engine == "sarvam":
            r = requests.post("https://api.sarvam.ai/speech-to-text", headers={"api-subscription-key": key}, timeout=10)
            if r.status_code in (401, 403):
                return False, "Invalid Sarvam AI API Key. Please check your credentials."
            return True, "Valid Sarvam AI API Key!"
        elif engine == "groq":
            r = requests.get("https://api.groq.com/openai/v1/models", headers={"Authorization": f"Bearer {key}"}, timeout=10)
            if r.status_code == 200:
                return True, "Valid Groq API Key!"
            return False, "Invalid Groq API Key. Please check your credentials."
        return True, "Valid."
    except Exception as e:
        return False, f"Verification failed: {e}"


# ----------------------------------------------------------------------------
# Ingestion
# ----------------------------------------------------------------------------

def download_from_gdrive(link: str, dest_dir: Path):
    """Download a Google Drive file/folder link (or direct URL). Returns media paths."""
    import gdown
    import inspect

    dest_dir.mkdir(parents=True, exist_ok=True)
    if "/folders/" in link:
        paths = gdown.download_folder(url=link, output=str(dest_dir), quiet=True) or []
    else:
        dl_kwargs = {"url": link, "output": str(dest_dir) + os.sep, "quiet": True}
        if "fuzzy" in inspect.signature(gdown.download).parameters:
            dl_kwargs["fuzzy"] = True
        out = gdown.download(**dl_kwargs)
        paths = [out] if out else []

    media = [p for p in paths if p and Path(p).suffix.lower() in MEDIA_EXTENSIONS]
    if not media:
        raise RuntimeError(
            "No downloadable media found at that link. Make sure it is shared as "
            "'Anyone with the link' and points to a video/audio file (or folder of them)."
        )
    return media


# ----------------------------------------------------------------------------
# Transcription engines
# ----------------------------------------------------------------------------

@st.cache_resource(show_spinner=False)
def load_whisper_model(model_size: str):
    from faster_whisper import WhisperModel
    try:
        return WhisperModel(model_size, device="cuda", compute_type="int8_float16"), "CUDA GPU"
    except Exception:
        return WhisperModel(model_size, device="cpu", compute_type="int8"), "CPU (int8)"


def transcribe_sarvam(audio_files, cfg, json_dir: Path, progress):
    """Batch transcription via the official sarvamai SDK (proven flow)."""
    import sarvamai

    client = sarvamai.SarvamAI(api_subscription_key=cfg["api_key"])
    stt_job = client.speech_to_text_job.create_job(
        model=cfg["sarvam_model"],
        language_code=cfg["language_code"],
        with_diarization=cfg["diarization"],
    )
    stt_job.upload_files(audio_files)
    stt_job.start()
    progress.progress(55, text="☁️ Sarvam AI cloud is transcribing...")

    while True:
        status = client.speech_to_text_job.get_status(stt_job.job_id)
        if status.job_state in ("Completed", "Failed"):
            break
        time.sleep(5)

    if status.job_state == "Failed":
        raise RuntimeError("Sarvam AI job failed on cloud servers. Check your API key/quota.")

    json_dir.mkdir(parents=True, exist_ok=True)
    stt_job.download_outputs(str(json_dir))
    json_files = sorted(glob.glob(str(json_dir / "*.json")))
    if not json_files:
        raise RuntimeError("Sarvam AI returned no transcript files.")
    return json_files


def transcribe_whisper(audio_files, model_size, json_dir: Path, progress):
    """100% local, offline transcription via faster-whisper."""
    model, device = load_whisper_model(model_size)
    st.caption(f"Faster-Whisper `{model_size}` on {device} — audio never leaves this machine.")
    json_dir.mkdir(parents=True, exist_ok=True)
    json_files = []
    for idx, a_file in enumerate(audio_files):
        progress.progress(50 + int(35 * idx / len(audio_files)),
                          text=f"🔒 Transcribing chunk {idx + 1}/{len(audio_files)} locally...")
        segments, _ = model.transcribe(a_file, beam_size=5)
        text = " ".join(seg.text.strip() for seg in segments)
        j_path = json_dir / f"chunk_{idx:02d}.json"
        with open(j_path, "w", encoding="utf-8") as f:
            json.dump({"transcript": text}, f, ensure_ascii=False)
        json_files.append(str(j_path))
    return json_files


def transcribe_groq(audio_files, api_key, json_dir: Path, progress):
    """100% Free & Ultra-Fast Cloud transcription via Groq API (whisper-large-v3)."""
    import requests

    if not api_key:
        raise RuntimeError("Groq API Key is required. Get a 100% free key at console.groq.com!")

    json_dir.mkdir(parents=True, exist_ok=True)
    json_files = []
    headers = {"Authorization": f"Bearer {api_key}"}
    url = "https://api.groq.com/openai/v1/audio/transcriptions"

    for idx, a_file in enumerate(audio_files):
        progress.progress(50 + int(35 * idx / len(audio_files)),
                          text=f"⚡ Groq Cloud transcribing chunk {idx + 1}/{len(audio_files)} at 100x speed...")
        
        with open(a_file, "rb") as f:
            files = {"file": (os.path.basename(a_file), f, "audio/wav")}
            data = {"model": "whisper-large-v3", "response_format": "json", "temperature": "0"}
            resp = requests.post(url, headers=headers, files=files, data=data, timeout=300)
            
        if resp.status_code != 200:
            raise RuntimeError(f"Groq API Error ({resp.status_code}): {resp.text}")
            
        res_json = resp.json()
        text = res_json.get("text", "").strip()
        
        j_path = json_dir / f"chunk_{idx:02d}.json"
        with open(j_path, "w", encoding="utf-8") as out_f:
            json.dump({"transcript": text}, out_f, ensure_ascii=False)
        json_files.append(str(j_path))
        
    return json_files


# ----------------------------------------------------------------------------
# Diarized document builder (speaker tags + timestamps)
# ----------------------------------------------------------------------------

def build_diarized_docx(json_files, output_path, source_name):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title_p = doc.add_paragraph()
    title_run = title_p.add_run(f"Transcript: {source_name}")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(99, 102, 241)
    title_p.paragraph_format.space_after = Pt(18)

    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run("Generated by VoxDoc AI Studio | High-Accuracy Diarized Transcript")
    meta_run.font.size = Pt(10)
    meta_run.font.italic = True
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    meta_p.paragraph_format.space_after = Pt(24)

    for path in sorted(json_files):
        with open(path, "r", encoding="utf-8") as f:
            chunk = json.load(f)

        entries = (chunk.get("diarized_transcript") or {}).get("entries", [])
        if not entries:
            # Fall back to plain transcript paragraphs
            for para in engine.split_into_paragraphs(chunk.get("transcript", ""), 6):
                p = doc.add_paragraph()
                r = p.add_run(para)
                r.font.size = Pt(11)
                p.paragraph_format.space_after = Pt(12)
            continue

        for entry in entries:
            speaker_id = entry.get("speaker_id", "Speaker")
            text = (entry.get("transcript") or "").strip()
            start_s = entry.get("start_time_seconds") or 0.0
            time_str = f"[{int(start_s // 60):02d}:{int(start_s % 60):02d}]"

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.line_spacing = 1.15
            spk_run = p.add_run(f"{speaker_id} {time_str}: ")
            spk_run.font.name = 'Arial'
            spk_run.font.bold = True
            spk_run.font.size = Pt(11)
            spk_run.font.color.rgb = RGBColor(79, 70, 229)
            txt_run = p.add_run(text)
            txt_run.font.size = Pt(11)

    doc.save(output_path)


# ----------------------------------------------------------------------------
# Full pipeline
# ----------------------------------------------------------------------------

def process_media(input_path: Path, cfg: dict, progress) -> dict:
    job_dir = input_path.parent
    audio_dir = job_dir / "audio"
    json_dir = job_dir / "json"
    docx_name = f"{input_path.stem}_VoxDoc_Transcript.docx"
    docx_path = job_dir / docx_name

    progress.progress(12, text="🎵 Step 1/3: Extracting resilient 16kHz WAV audio chunks...")
    audio_files, duration = engine.extract_and_segment_audio(
        str(input_path), str(audio_dir), chunk_minutes=cfg["chunk_minutes"])

    progress.progress(40, text=f"🤖 Step 2/3: {len(audio_files)} chunk(s) ready "
                               f"({duration / 60:.1f} min). Transcribing...")

    if cfg["engine"] == "sarvam":
        json_files = transcribe_sarvam(audio_files, cfg, json_dir, progress)
        model_label = f"Sarvam AI {cfg['sarvam_model']} Batch Speech-to-Text"
    elif cfg["engine"] == "groq":
        json_files = transcribe_groq(audio_files, cfg["groq_api_key"], json_dir, progress)
        model_label = "Groq Cloud Whisper-Large-V3 (100% Free & Ultra-Fast)"
    else:
        json_files = transcribe_whisper(audio_files, cfg["whisper_model"], json_dir, progress)
        model_label = f"Faster-Whisper {cfg['whisper_model']} (Local, Offline)"

    progress.progress(90, text="📄 Step 3/3: Formatting professional Word document...")
    if cfg["engine"] == "sarvam" and cfg["diarization"]:
        build_diarized_docx(json_files, str(docx_path), input_path.name)
    else:
        engine.create_transcription_docx(
            json_files, str(docx_path), input_path.name, duration,
            font_name=cfg["font_name"],
            sentences_per_paragraph=cfg["para_sentences"],
            model_label=model_label,
            chunk_minutes=cfg["chunk_minutes"],
        )
    progress.progress(100, text="✨ Transcription Complete!")

    words, raw_chunks = 0, []
    for jf in json_files:
        with open(jf, "r", encoding="utf-8") as f:
            data = json.load(f)
        raw_chunks.append(data)
        words += len((data.get("transcript") or "").split())

    return {
        "source": input_path.name,
        "docx_name": docx_name,
        "docx_bytes": docx_path.read_bytes(),
        "json_text": json.dumps(raw_chunks, indent=2, ensure_ascii=False),
        "duration_min": duration / 60,
        "chunks": len(audio_files),
        "words": words,
        "model_label": model_label,
    }


def render_results():
    results = st.session_state.get("results", [])
    if not results:
        return
    st.markdown("---")
    st.markdown("### 🎉 Completed Transcripts — Ready for Download!")
    for i, r in enumerate(results):
        with st.container(border=True):
            st.markdown(f"**📄 {r['docx_name']}**")
            c1, c2, c3 = st.columns(3)
            c1.metric("Duration", f"{r['duration_min']:.1f} min")
            c2.metric("Chunks", r["chunks"])
            c3.metric("Words", f"{r['words']:,}")
            st.caption(f"Source: `{r['source']}` · Engine: {r['model_label']}")

            d1, d2 = st.columns(2)
            with d1:
                st.download_button(
                    "📥 Download Word Doc (.docx)",
                    data=r["docx_bytes"],
                    file_name=r["docx_name"],
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_docx_{i}",
                    use_container_width=True,
                )
            with d2:
                st.download_button(
                    "📥 Download Raw JSON Data",
                    data=r["json_text"],
                    file_name=r["docx_name"].replace(".docx", ".json"),
                    mime="application/json",
                    key=f"dl_json_{i}",
                    use_container_width=True,
                )

            with st.expander("👁️ Preview Transcript Text"):
                for chunk in json.loads(r["json_text"]):
                    entries = (chunk.get("diarized_transcript") or {}).get("entries", [])
                    if entries:
                        for entry in entries:
                            st.markdown(f"**{entry.get('speaker_id', 'Speaker')}:** "
                                        f"{entry.get('transcript', '')}")
                    else:
                        st.write((chunk.get("transcript") or "")[:3000])


# ----------------------------------------------------------------------------
# Main UI
# ----------------------------------------------------------------------------

def main():
    if not check_password():
        return

    st.markdown("<div class='main-header'>🎙️ VoxDoc AI Studio <span style='font-size: 1.1rem; background: #10B981; color: white; padding: 4px 12px; border-radius: 20px; vertical-align: middle;'>🔒 Private Cloud Mode</span></div>", unsafe_allow_html=True)
    st.markdown("<div class='sub-header'>Turn any video, recording, or Google Drive link into "
                "beautifully formatted Word documents in minutes — powered by Sarvam AI, Groq & "
                "Faster-Whisper.</div>", unsafe_allow_html=True)

    # --- Sidebar settings ---
    with st.sidebar:
        st.markdown("### 🔒 Personal Workspace")
        st.success("Authenticated as **Master Admin**")
        if st.button("🚪 Lock Workspace / Log Out", use_container_width=True):
            st.session_state["authenticated"] = False
            st.rerun()
        st.markdown("---")

        st.markdown("### ⚙️ Engine Settings")

        ai_engine = st.selectbox(
            "Select AI Transcription Engine",
            [
                "Sarvam AI Batch API (Best for Tamil/Hindi/Indian)",
                "Groq Cloud Whisper-Large-V3 (100% Free & Ultra-Fast)",
                "Local Faster-Whisper (Offline & Private)"
            ],
        )
        if "Sarvam" in ai_engine:
            engine_key = "sarvam"
        elif "Groq" in ai_engine:
            engine_key = "groq"
        else:
            engine_key = "whisper"

        api_key = get_secret("SARVAM_API_KEY", "")
        groq_api_key = get_secret("GROQ_API_KEY", "")
        language_code = "en-IN"
        sarvam_model = "saaras:v3"
        whisper_model = "large-v3"
        diarization = False
        api_key_valid = True

        if engine_key == "sarvam":
            api_key = st.text_input(
                "Sarvam API Key", value=api_key, type="password",
                help="Get a key at dashboard.sarvam.ai — or set SARVAM_API_KEY in secrets.",
            )
            if api_key.strip():
                is_valid, msg = verify_api_key("sarvam", api_key)
                if is_valid:
                    st.success("✅ **Verified & Valid!** Ready to transcribe.")
                    api_key_valid = True
                else:
                    st.error(f"❌ **Invalid Key!** {msg}")
                    api_key_valid = False
            else:
                st.caption("⚠️ Please enter a Sarvam API Key to continue.")
                api_key_valid = False

            language_code = LANGUAGE_OPTIONS[st.selectbox("Audio Language", list(LANGUAGE_OPTIONS))]
            sarvam_model = SARVAM_MODELS[st.selectbox("Sarvam Model", list(SARVAM_MODELS))]
            diarization = st.toggle(
                "🗣️ Speaker Diarization",
                value=False,
                help="Label speakers (Speaker 1, Speaker 2...) with timestamps in the Word document.",
            )
        elif engine_key == "groq":
            groq_api_key = st.text_input(
                "Groq API Key (100% Free)", value=groq_api_key, type="password",
                help="Get a free key instantly at console.groq.com — or set GROQ_API_KEY in secrets.",
            )
            if groq_api_key.strip():
                is_valid, msg = verify_api_key("groq", groq_api_key)
                if is_valid:
                    st.success("✅ **Verified & Valid!** Ready to transcribe.")
                    api_key_valid = True
                else:
                    st.error(f"❌ **Invalid Key!** {msg}")
                    api_key_valid = False
            else:
                st.caption("⚠️ Please enter a Groq API Key to continue.")
                api_key_valid = False

            st.info("⚡ Groq provides 100% FREE ultra-fast cloud transcription using Whisper-Large-V3. No GPU required on hosting server!")
        else:
            whisper_model = st.selectbox(
                "Whisper Model Size", ["large-v3", "medium", "small", "base", "tiny"], index=0,
                help="large-v3 = best accuracy. Uses your GPU when CUDA is available.",
            )
            api_key_valid = True

        chunk_minutes = st.slider(
            "Audio Chunk Size (Minutes)", 15, 45, 40,
            help="Splits large recordings to prevent timeout/size limits.",
        )

        st.markdown("---")
        st.markdown("### 📑 Word Document Styling")
        font_name = FONT_OPTIONS[st.selectbox("Font Family", list(FONT_OPTIONS))]
        para_sentences = st.slider("Sentences per Paragraph", 3, 10, 6)

        st.markdown("---")
        st.markdown("#### 🚀 About VoxDoc AI")
        st.caption("Resilient FFmpeg audio processing, Sarvam AI, Groq Cloud & local "
                   "Whisper — configured for private cloud deployment (`naa matum use pannura mathuri`).")

    cfg = {
        "engine": engine_key,
        "api_key": api_key,
        "groq_api_key": groq_api_key,
        "api_key_valid": api_key_valid,
        "language_code": language_code,
        "sarvam_model": sarvam_model,
        "whisper_model": whisper_model,
        "diarization": diarization,
        "chunk_minutes": chunk_minutes,
        "font_name": font_name,
        "para_sentences": para_sentences,
    }

    # --- Ingestion tabs ---
    tab1, tab2 = st.tabs(["📁 Upload Local File", "🔗 Google Drive Import"])

    with tab1:
        uploaded_file = st.file_uploader(
            "Drop your video or audio recording here",
            type=[e.lstrip(".") for e in sorted(MEDIA_EXTENSIONS)],
        )
        if uploaded_file is not None:
            st.caption(f"**{uploaded_file.name}** · {uploaded_file.size / 1048576:.1f} MB")

    with tab2:
        gdrive_link = st.text_input(
            "Paste Google Drive Shareable Link (file or folder)",
            placeholder="https://drive.google.com/file/d/1XyZ.../view?usp=sharing",
            help="The link must be shared as 'Anyone with the link can view'.",
        )

    # --- Action button & pipeline execution ---
    if st.button("🚀 Start Transcription Pipeline", use_container_width=True, type="primary"):
        if cfg["engine"] == "sarvam" and not cfg["api_key"].strip():
            st.error("⚠️ Please enter your Sarvam AI API key in the sidebar "
                     "(or switch to the free Groq or Local Faster-Whisper engine).")
        elif cfg["engine"] == "sarvam" and not cfg.get("api_key_valid", False):
            st.error("❌ Your Sarvam AI API key is invalid! Please verify and enter a valid API key in the sidebar.")
        elif cfg["engine"] == "groq" and not cfg["groq_api_key"].strip():
            st.error("⚠️ Please enter your Groq API key in the sidebar (get a 100% free key at console.groq.com!).")
        elif cfg["engine"] == "groq" and not cfg.get("api_key_valid", False):
            st.error("❌ Your Groq API key is invalid! Please verify and enter a valid API key in the sidebar.")
        elif uploaded_file is None and not gdrive_link.strip():
            st.warning("⚠️ Please upload a media file or paste a Google Drive link first!")
        else:
            st.markdown("### ⚡ Pipeline Progress")
            progress = st.progress(3, text="Initializing pipeline...")
            job_dir = WORKSPACE_ROOT / f"job_{uuid.uuid4().hex[:8]}"
            job_dir.mkdir(parents=True, exist_ok=True)
            try:
                if uploaded_file is not None:
                    input_path = job_dir / Path(uploaded_file.name).name
                    with open(input_path, "wb") as f:
                        f.write(uploaded_file.getbuffer())
                    media_files = [str(input_path)]
                    st.success(f"Loaded local file: `{uploaded_file.name}`")
                else:
                    if not re.search(r"https?://", gdrive_link):
                        raise RuntimeError("Please paste a valid link (https://drive.google.com/...).")
                    progress.progress(6, text="⬇️ Downloading from Google Drive...")
                    media_files = download_from_gdrive(gdrive_link.strip(), job_dir)
                    st.success(f"Downloaded {len(media_files)} media file(s) from Drive.")

                for m in media_files:
                    result = process_media(Path(m), cfg, progress)
                    st.session_state.setdefault("results", []).insert(0, result)
                st.balloons()
            except Exception as e:
                st.error(f"❌ Pipeline failed: {e}")

    render_results()


if __name__ == "__main__":
    main()
