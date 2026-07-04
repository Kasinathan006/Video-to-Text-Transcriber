"""
Transcribe a folder of WhatsApp voice notes (.ogg) into one clean,
plain white Word document.

SETUP (run once):
    pip install openai-whisper python-docx
    # ffmpeg must also be installed and on PATH:
    #   Windows: https://ffmpeg.org/download.html (add to PATH)
    #   Mac:     brew install ffmpeg
    #   Linux:   sudo apt install ffmpeg

USAGE:
    1. Put all your .ogg files in the audio_files folder (or edit AUDIO_DIR below).
    2. Run: python transcribe_to_docx.py
    3. Open transcript.docx
"""

import os
import glob
import whisper
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ---------------- SETTINGS ----------------
AUDIO_DIR = "."                # folder containing your .ogg files (default: current directory)
OUTPUT_DOCX = "transcript.docx"
MODEL_SIZE = "large-v3"        # best accuracy for Tamil/English mixed speech
LANGUAGE = None                # None = auto-detect per file, or set "ta" to force Tamil
# --------------------------------------------


def get_audio_files(folder):
    files = []
    for ext in ("*.wav", "*.ogg", "*.mp3", "*.m4a"):
        files.extend(glob.glob(os.path.join(folder, ext)))
    files = sorted(set(files))
    if not files:
        raise FileNotFoundError(f"No audio files (.wav, .ogg, .mp3, .m4a) found in '{folder}'")
    return files


def transcribe_file(model, path, language):
    result = model.transcribe(path, language=language, task="transcribe", verbose=False)
    return result.get("text", "").strip()


def build_docx(entries, output_path):
    doc = Document()

    # Configure style with Unicode / Complex Script support (prevents Tamil "box" formatting issues)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    
    # Explicitly set Complex Script (cs) font for Tamil/Tanglish rendering in Word
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:cs"), "Nirmala UI")  # Nirmala UI natively supports Tamil on Windows

    title = doc.add_heading("Voice Notes Transcript", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for filename, text in entries:
        doc.add_heading(filename, level=2)
        if text:
            doc.add_paragraph(text)
        else:
            doc.add_paragraph("[No speech detected]")

    doc.save(output_path)


def main():
    files = get_audio_files(AUDIO_DIR)
    print(f"Found {len(files)} audio files. Loading Whisper model '{MODEL_SIZE}'...")
    model = whisper.load_model(MODEL_SIZE)

    entries = []
    for i, path in enumerate(files, start=1):
        filename = os.path.basename(path)
        print(f"[{i}/{len(files)}] Transcribing {filename} ...")
        text = transcribe_file(model, path, LANGUAGE)
        entries.append((filename, text))
        print(f"    -> {text[:80]}{'...' if len(text) > 80 else ''}")

    build_docx(entries, OUTPUT_DOCX)
    print(f"\nDone. Full transcript saved to: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
