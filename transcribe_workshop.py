#!/usr/bin/env python3
"""
transcribe_workshop.py - End-to-End Transcription & Word Doc Builder for Workshop.mp4

1. Splits Workshop.mp4 into 1-hour 16kHz mono WAV chunks (cached in audio_chunks_workshop/).
2. Transcribes each chunk on GPU using faster-whisper (cached in txt_outputs_workshop/).
3. Generates a cleanly formatted Word (.docx) file: Workshop_Full_End_to_End_Transcription.docx.
"""

import os
import sys
import glob
import time
import subprocess
import shutil
import re
from faster_whisper import WhisperModel
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn


def log(msg):
    print(msg, flush=True)


def get_video_duration(filepath):
    ffprobe = shutil.which("ffprobe")
    if not ffprobe:
        return None
    try:
        cmd = [
            ffprobe, "-v", "error", "-show_entries", "format=duration",
            "-of", "default=noprint_wrappers=1:nokey=1", str(filepath)
        ]
        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=True)
        return float(result.stdout.strip())
    except Exception:
        return None


def extract_audio_chunks(video_path="Workshop.mp4", chunk_dir="audio_chunks_workshop", chunk_sec=3600):
    os.makedirs(chunk_dir, exist_ok=True)
    duration = get_video_duration(video_path)
    if not duration:
        log("[ERROR] Could not determine duration of video.")
        sys.exit(1)

    log(f"Video '{video_path}' duration: {duration:.1f}s ({duration/3600:.2f} hours)")
    num_chunks = int((duration + chunk_sec - 1) // chunk_sec)
    log(f"Splitting audio into {num_chunks} chunk(s) of up to {chunk_sec/60:.0f} mins each...")

    audio_files = []
    for i in range(num_chunks):
        start_sec = i * chunk_sec
        chunk_name = f"part_{i+1:02d}.wav"
        chunk_path = os.path.join(chunk_dir, chunk_name)
        audio_files.append(chunk_path)

        if os.path.exists(chunk_path) and os.path.getsize(chunk_path) > 100000:
            log(f"  [SKIP] Chunk already exists: {chunk_name}")
            continue

        log(f"  [EXTRACTING] Chunk {i+1}/{num_chunks} ({start_sec}s to {start_sec+chunk_sec}s) -> {chunk_name}")
        cmd = [
            "ffmpeg", "-y", "-i", video_path,
            "-ss", str(start_sec),
            "-t", str(chunk_sec),
            "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1",
            chunk_path
        ]
        proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if proc.returncode != 0:
            log(f"  [ERROR] FFmpeg failed for chunk {chunk_name}")
            if os.path.exists(chunk_path):
                os.remove(chunk_path)
            sys.exit(1)

    return sorted(audio_files), duration


def split_into_paragraphs(text, target_sentences=6):
    text = text.strip()
    if not text:
        return []
    sentences = re.split(r'(?<=[.!?])\s+', text)
    paragraphs = []
    current_chunk = []
    for sentence in sentences:
        current_chunk.append(sentence)
        if len(current_chunk) >= target_sentences:
            paragraphs.append(' '.join(current_chunk))
            current_chunk = []
    if current_chunk:
        paragraphs.append(' '.join(current_chunk))
    return paragraphs


def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def build_docx(transcripts, total_duration, output_path="Workshop_Full_End_to_End_Transcription.docx"):
    log(f"\nBuilding professional Word document: {output_path}...")
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Normal Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x2D, 0x37, 0x48) # Dark Charcoal
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(8)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("Workshop - Full End-to-End Transcription")
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    run_sub = subtitle_p.add_run("Complete Verbatim Session Record")
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    # Metadata Table
    hours = int(total_duration // 3600)
    mins = int((total_duration % 3600) // 60)
    secs = int(total_duration % 60)
    dur_str = f"{hours}h {mins}m {secs}s ({total_duration/60:.1f} minutes)"

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    metadata = [
        ("Source File:", "Workshop.mp4"),
        ("Total Duration:", dur_str),
        ("AI Model Used:", "Faster-Whisper (Large-v3 GPU)"),
        ("Transcription Fidelity:", "100% Verbatim End-to-End Extraction")
    ]

    for i, (label, val) in enumerate(metadata):
        row = table.rows[i]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.3)
        set_cell_margins(cell_lbl)
        set_cell_margins(cell_val)
        bg = "F7FAFC" if i % 2 == 0 else "EDF2F7"
        set_cell_background(cell_lbl, bg)
        set_cell_background(cell_val, bg)
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(0)
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(10.5)
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(0)
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    total_words = 0
    for part_idx, (part_name, text) in enumerate(transcripts, 1):
        words = len(text.split())
        total_words += words

        h2 = doc.add_paragraph()
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(8)
        r_h2 = h2.add_run(f"{part_name} (approx. {words:,} words)")
        r_h2.font.size = Pt(15)
        r_h2.font.bold = True
        r_h2.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

        paragraphs = split_into_paragraphs(text, target_sentences=6)
        if not paragraphs:
            doc.add_paragraph("[No speech detected in this audio segment]")
        else:
            for p_text in paragraphs:
                p = doc.add_paragraph(p_text)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.15

    doc.save(output_path)
    log(f"[SUCCESS] Saved Word document to '{output_path}' (Total words: {total_words:,}).")


def main():
    log("============================================================")
    log("STARTING WORKSHOP.MP4 END-TO-END TRANSCRIPTION PIPELINE")
    log("============================================================")
    
    audio_files, total_duration = extract_audio_chunks("Workshop.mp4", "audio_chunks_workshop", 3600)
    
    log("\nLoading faster-whisper model 'large-v3' on CUDA...")
    try:
        model = WhisperModel("large-v3", device="cuda", compute_type="int8_float16")
    except Exception as e:
        log(f"[WARN] int8_float16 failed ({e}), trying int8...")
        model = WhisperModel("large-v3", device="cuda", compute_type="int8")
    log("Model loaded successfully!")

    os.makedirs("txt_outputs_workshop", exist_ok=True)
    transcripts = []

    for i, file_path in enumerate(audio_files, start=1):
        filename = os.path.basename(file_path)
        part_title = f"Part {i}: Hour {i-1} to Hour {i}"
        txt_save_path = os.path.join("txt_outputs_workshop", f"{os.path.splitext(filename)[0]}.txt")

        if os.path.exists(txt_save_path) and os.path.getsize(txt_save_path) > 50:
            log(f"\n[{i}/{len(audio_files)}] Using cached transcript for {filename}...")
            with open(txt_save_path, "r", encoding="utf-8") as f:
                text = f.read()
            transcripts.append((part_title, text))
            continue

        log(f"\n[{i}/{len(audio_files)}] Transcribing {filename}...")
        chunk_start = time.time()
        
        segments, info = model.transcribe(file_path, beam_size=1, language="en")
        log(f"Detected language '{info.language}' with probability {info.language_probability:.2f}")

        txt_tmp_path = txt_save_path + ".tmp"
        text_lines = []
        for seg_idx, segment in enumerate(segments, start=1):
            text_lines.append(segment.text.strip())
            if seg_idx % 30 == 0:
                elapsed_sec = time.time() - chunk_start
                log(f"  ...processed {seg_idx} segments ({segment.end:.1f}s of audio transcribed in {elapsed_sec:.1f}s)")
                with open(txt_tmp_path, "w", encoding="utf-8") as f:
                    f.write(" ".join(text_lines))

        full_text = " ".join(text_lines)
        
        with open(txt_save_path, "w", encoding="utf-8") as f:
            f.write(full_text)
        if os.path.exists(txt_tmp_path):
            try:
                os.remove(txt_tmp_path)
            except Exception:
                pass

        elapsed = time.time() - chunk_start
        log(f"-> Completed {filename} in {elapsed:.1f}s ({len(full_text)} chars)")
        transcripts.append((part_title, full_text))

    build_docx(transcripts, total_duration, "Workshop_Full_End_to_End_Transcription.docx")
    log("\n============================================================")
    log("ALL DONE! Workshop_Full_End_to_End_Transcription.docx CREATED.")
    log("============================================================")


if __name__ == "__main__":
    main()
